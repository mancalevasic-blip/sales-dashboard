"""Generate a Word document summarizing the sales analytics project end to end."""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATA_PATH = "data/orders.csv"
OUTPUT_PATH = "project_summary.docx"


def load_orders():
    df = pd.read_csv(DATA_PATH)
    df["order_date"] = pd.to_datetime(df["order_date"])
    df["revenue"] = df["quantity"] * df["unit_price"]
    df["month"] = df["order_date"].dt.strftime("%Y-%m")
    return df


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build():
    df = load_orders()
    monthly = (
        df.groupby("month")
        .agg(revenue=("revenue", "sum"), orders=("order_id", "nunique"))
        .reset_index()
        .sort_values("month")
    )
    by_cat = (
        df.groupby("category")
        .agg(revenue=("revenue", "sum"), orders=("order_id", "nunique"))
        .reset_index()
    )
    by_cat["aov"] = by_cat["revenue"] / by_cat["orders"]
    by_cat = by_cat.sort_values("revenue", ascending=False)

    total_revenue = df["revenue"].sum()
    total_orders = df["order_id"].nunique()
    aov = total_revenue / total_orders
    best = monthly.loc[monthly["revenue"].idxmax()]
    h2_2024 = monthly.loc[monthly["month"].between("2024-07", "2024-12"), "revenue"].sum()
    h1_2025 = monthly.loc[monthly["month"].between("2025-01", "2025-06"), "revenue"].sum()
    growth = (h1_2025 / h2_2024 - 1) * 100

    doc = Document()

    title = doc.add_heading("Sales Analytics Project — Summary", level=0)

    doc.add_paragraph(
        "This document summarizes the sales analytics work delivered for this project: "
        "the source data, the analysis and its conclusions, and the deliverables built "
        "for management (a static report, a live dashboard, and a board-level slide deck)."
    )

    add_heading(doc, "1. Data Source")
    doc.add_paragraph(
        f"All analysis is based on data/orders.csv — {len(df):,} order line items "
        f"covering {monthly['month'].min()} to {monthly['month'].max()}, "
        f"{df['category'].nunique()} product categories, {df['product_id'].nunique()} products, "
        f"and {df['customer_id'].nunique():,} unique customers."
    )
    add_bullets(
        doc,
        [
            "Columns: order_id, order_date, customer_id, product_id, product_name, category, quantity, unit_price",
            "Revenue is calculated as quantity × unit_price (gross revenue)",
            "No discount, return, or refund columns are present — figures are not net of returns",
            "This is a one-time snapshot; there is no recurring data refresh planned",
        ],
    )

    add_heading(doc, "2. Key Findings")
    add_bullets(
        doc,
        [
            f"Total revenue: ${total_revenue:,.2f} across {total_orders:,} orders",
            f"Average order value: ${aov:,.2f}, stable within a narrow band all year",
            f"Best month: {best['month']} with ${best['revenue']:,.2f} in revenue ({int(best['orders'])} orders)",
            "December's peak is driven by order volume, not by customers spending more per order "
            "(holiday seasonality, not a pricing effect)",
            "January shows the only double-digit month-over-month revenue decline, followed by a "
            "steady five-month recovery",
            f"Underlying (non-holiday) demand is growing: H2 2024 revenue was ${h2_2024:,.2f} vs. "
            f"${h1_2025:,.2f} in H1 2025, a {growth:+.1f}% change despite H1 2025 having no holiday spike",
        ],
    )

    add_heading(doc, "3. Revenue by Month")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Month", "Revenue", "Orders"
    for row in monthly.itertuples():
        cells = table.add_row().cells
        cells[0].text = row.month
        cells[1].text = f"${row.revenue:,.2f}"
        cells[2].text = f"{row.orders:,}"

    add_heading(doc, "4. Revenue by Category")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Category", "Revenue", "Orders", "AOV"
    for row in by_cat.itertuples():
        cells = table.add_row().cells
        cells[0].text = row.category
        cells[1].text = f"${row.revenue:,.2f}"
        cells[2].text = f"{row.orders:,}"
        cells[3].text = f"${row.aov:,.2f}"

    add_heading(doc, "5. Deliverables")

    add_heading(doc, "5.1 Static analysis report", level=2)
    doc.add_paragraph(
        "monthly_sales_report.qmd — a Quarto/R report with the original revenue, AOV, and "
        "seasonality analysis, rendered to monthly_sales_report.html and monthly_sales_report.pdf."
    )

    add_heading(doc, "5.2 Interactive dashboard", level=2)
    doc.add_paragraph(
        "streamlit_app.py — a Streamlit dashboard with KPI cards, a revenue trend chart, a "
        "category breakdown, and sidebar filters for category and date range. Code is hosted at "
        "https://github.com/mancalevasic-blip/sales-dashboard, deployable to Streamlit Community "
        "Cloud as a public, no-login app. Requirements: docs/brainstorms/2026-06-22-sales-dashboard-requirements.md"
    )

    add_heading(doc, "5.3 Leadership/board deck", level=2)
    doc.add_paragraph(
        "build_sales_deck.py generates sales_board_deck.pptx — an 8-slide deck (title, KPI "
        "summary, revenue trend, seasonality narrative, category breakdown, category detail, "
        "underlying growth, takeaways) in a light-editorial visual style. Requirements: "
        "docs/brainstorms/2026-06-22-sales-board-deck-requirements.md"
    )

    add_heading(doc, "6. Limitations")
    add_bullets(
        doc,
        [
            "Revenue figures are gross, with no adjustment for returns, discounts, or refunds",
            "Only twelve months of data are available, so the seasonal interpretation of the "
            "December peak cannot be confirmed without a second holiday cycle to compare against",
            "Drill-down is limited to category and date — no product- or customer-level breakdown "
            "was requested or built",
        ],
    )

    doc.save(OUTPUT_PATH)
    print(f"Saved {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
